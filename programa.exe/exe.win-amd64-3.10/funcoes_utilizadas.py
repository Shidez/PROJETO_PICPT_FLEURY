import openpyxl
import shutil
import pandas as pd
import os
import datetime
from docx import Document
from docx.shared import Pt


def backup_excel(caminho_arquivo):
    nome_backup = os.path.basename(caminho_arquivo).replace(".xlsm", "_BACKUP.xlsm")
    diretorio_backup = os.path.join(os.path.dirname(caminho_arquivo), "BACKUP")
    caminho_backup = os.path.join(diretorio_backup, nome_backup)
  

    if not os.path.isdir(diretorio_backup):
        os.mkdir(diretorio_backup)
        print("pasta criada")
    else:
        print("pasta ja existente")

    shutil.copy2(caminho_arquivo, caminho_backup)
        
        

    

    



def baixa_postal(caminho_arquivo, pdfs_encontrados):
    # lista criada na função copia_pdfs
    pdfs_encontrados = list(map(int, pdfs_encontrados))

    workbook = openpyxl.load_workbook(
        filename=caminho_arquivo, read_only=False, keep_vba=True)

    worksheet = workbook['DADOS']
    # percorrer as células da coluna C
    for cell in worksheet['C']:
        # verificar se o valor da célula está no set
        if cell.value in pdfs_encontrados:
            # inserir 'OK' na coluna T da mesma linha
            worksheet.cell(row=cell.row, column=20).value = 'OK'
    # salvar as alterações na planilha
    workbook.save(caminho_arquivo)


def fichas_nome(caminho_arquivo, pdfs_encontrados):
    try:
        # Criar uma lista para guardar os números das linhas como global
        global fichas_nomes
        fichas_nomes = {}
        # Ler o arquivo Excel
        df = pd.read_excel(
            caminho_arquivo, sheet_name="DADOS", engine="openpyxl")
        # Acessar a coluna R
        coluna_r = df["Ficha"]

        for indice, valor in coluna_r.items():
            for item in pdfs_encontrados:
                if valor == int(item):
                    ficha = df.at[indice, "Ficha"]
                    ficha = int(ficha)
                    nome = df.at[indice, "Nome"]
                    fichas_nomes.update({str(ficha): nome})

        return fichas_nomes

    except Exception as erro:
        print('ERRO!', f"{erro}")


def escreve_arquivo(caminho, fichas_nomes, caminho_arquivo):

    data_atual = datetime.date.today().strftime("%d-%m-%Y")
    os.chdir(caminho)

    documento = Document()
    
    if "ONCOTYPE" in caminho_arquivo.upper():
        texto = "CONFERÊNCIA DE FICHAS ONCOTYPE"
    elif "MIRTHYPE" in caminho_arquivo.upper():
        texto = "CONFERÊNCIA DE FICHAS MIRTHYPE"
    else:
        texto = "CONFERÊNCIA DE FICHAS FOUNDATION"

    # Adicionar o título no documento
    
    texto_centralizado = texto.center(90, '-')
    p = documento.add_paragraph()
    run = p.add_run(texto_centralizado)
    run.bold = True
    run.font.size = Pt(12)

    # Criar a tabela
    tabela = documento.add_table(rows=1, cols=4)

    # Definir o estilo da tabela
    tabela.style = "Table Grid"

    # Definir a largura das colunas
    for i in range(4):
        # Definindo a largura como 60 pontos para cada coluna
        tabela.columns[i].width = Pt(60)

    # Adicionar cabeçalho
    cabecalho = tabela.rows[0].cells
    cabecalho[0].text = "Ficha"
    cabecalho[1].text = "Nome"
    cabecalho[2].text = "Conferência 1"
    cabecalho[3].text = "Conferência 2"
    # cabecalho[4].text = data_atual

    # Estilizar cabeçalho em negrito
    for cell in cabecalho:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Adicionar as fichas e nomes na tabela
    for chave, valor in fichas_nomes.items():
        linha = tabela.add_row().cells
        # tem que ser em str o número para salvar em word ou txt
        linha[0].text = str(chave)
        linha[1].text = valor
        linha[2].text = ""
        linha[3].text = ""

    # Salvar o documento no formato .docx
    documento.save(f"{data_atual}.docx")


def escreve_arquivo_não_encontrado(caminho, lista):
    os.chdir(caminho)
    texto = "Fichas não encontradas"

    texto_centralizado = texto.center(100, '-')

    if os.path.isfile(caminho + f"/Fichas não encontradas.txt"):
        with open(f"Fichas não encontradas.txt", "a") as f:
            for numero in lista:
                f.write(f"{str(numero)}" + "\n")

    else:
        with open(f"Fichas não encontradas.txt", "w") as f:
            f.write(f"{texto_centralizado}" + "\n\n")
            for numero in lista:
                f.write(str(numero) + "\n")



